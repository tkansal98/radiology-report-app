import streamlit as st
import json
from docx import Document
from io import BytesIO

@st.cache_data
def load_templates():
    with open("templates.json", "r", encoding="utf-8") as f:
        return json.load(f)

templates = load_templates()

modality = st.selectbox("Select Modality", sorted(set(t["modality"] for t in templates)))
system = st.selectbox("Select System", sorted(set(t["system"] for t in templates if t["modality"] == modality)))
filtered = [t for t in templates if t["modality"] == modality and t["system"] == system]
diagnosis = st.selectbox("Select Diagnosis", ["Diagnosis Unknown"] + sorted(set(t["diagnosis"] for t in filtered)))

template = None
if diagnosis == "Diagnosis Unknown":
    template = next((t for t in filtered if "unknown" in t["diagnosis"].lower()), None)
else:
    template = next((t for t in filtered if t["diagnosis"] == diagnosis), None)

if template:
    st.subheader("Fill Report Fields")
    st.info("💡 Tip: Click in any text box and press **Win + H** to enable Windows voice dictation.")

    report_data = {}
    for field_spec in template["inputs"]:
        field = field_spec["field"]
        label = field_spec.get("label", field.replace("_", " ").title())
        normal = field_spec.get("normal", "")
        report_data[field] = st.text_area(label, key=field, placeholder=normal)

    if st.button("Download Report as .docx"):
        doc = Document()
        doc.add_heading(f"{modality} {system} – {diagnosis}", level=1)


        buffer = BytesIO()
        for field in template["inputs"]:
            if label.startswith('Frontal - '): label = label.replace('Frontal - ', '')
            if label.startswith('Lateral - '): label = label.replace('Lateral - ', '')
            label = field.get("label", field["field"])
            val = report_data.get(field["field"], "").strip()
            if not val:
                val = field.get("normal", "")
            if val:
                para = doc.add_paragraph()
                run = para.add_run(f"{label}: ")
                if val != field.get("normal", ""):
                    val_run = para.add_run(val)
                    val_run.bold = True
                else:
                    para.add_run(val)
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Download .docx", buffer, file_name="Radiology_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.warning("No template found.")